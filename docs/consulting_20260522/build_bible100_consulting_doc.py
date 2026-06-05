from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parent / "Bible100_網站改良與發展建議_2026-05-22.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def style_run(run, size=None, bold=None, color=None):
    run.font.name = "Microsoft JhengHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc, text, level=1):
    p = doc.add_heading("", level=level)
    r = p.add_run(text)
    style_run(r, size={1: 16, 2: 13, 3: 12}.get(level, 12), bold=True, color="2E74B5" if level < 3 else "1F4D78")
    return p


def add_para(doc, text="", bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        style_run(r1, size=11, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        style_run(r2, size=11)
    else:
        r = p.add_run(text)
        style_run(r, size=11)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        style_run(r, size=11)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.1


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        style_run(r, size=11)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.1


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = Inches(widths[i])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "F2F4F7")
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        style_run(r, size=10, bold=True, color="1F4D78")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].width = Inches(widths[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[i])
            p = cells[i].paragraphs[0]
            r = p.add_run(val)
            style_run(r, size=10)
    doc.add_paragraph()
    return table


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft JhengHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    normal.font.size = Pt(11)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = title.add_run("Bible100 網站改良與發展建議")
    style_run(r, size=24, bold=True, color="0B2545")
    sub = doc.add_paragraph()
    r = sub.add_run("給非技術維護者、山區多語聖經老師與教會事工同工使用的初步顧問報告 | 2026-05-22")
    style_run(r, size=10, color="555555")

    add_heading(doc, "一、我看到的網站定位", 1)
    add_para(doc, "Bible100 不是普通首頁式網站，而是一個多模組、可拆成多個獨立小站的聖經學習與教會事工平台。index_v5.html 是總入口，負責把教材、聖經研讀、AI 工具、教會事工、學校管理、詩歌管理等模組集中起來。")
    add_para(doc, "教育角度：它的核心不是展示，而是幫助老師在不同語言、不同網絡條件、不同電腦能力下，把聖經教學和教會行政工作做穩。")

    add_heading(doc, "二、目前資料夾與模組觀察", 1)
    add_table(
        doc,
        ["範圍", "目前觀察", "改良方向"],
        [
            ["總入口", "index.html 會轉向 index_v5.html；index_v5 是雙 iframe 總站殼。", "保留總入口，強化導覽、錯誤提示、手機版與維護說明。"],
            ["主要模組", "bible_study、ai_tools、church_ministry、church_planning、school_management、languages、hymn_management 等。", "每個模組整理成可獨立部署的小站，統一入口、README、測試清單。"],
            ["測試與審計", "已有 tests/run-all-tests.ps1、site_full_audit.py、斷鏈檢查等工具。", "先修準測試工具，再建立每次修改後的固定檢查流程。"],
            ["文件", "docs 內已有路線圖、改良計劃、資料契約、維護文件。", "把工程文件整理成小白版、老師版、開發者版三層。"],
            ["AI 與同步", "index_v5 已有 AI 工具入口與同步紀錄抽屜。", "逐步加入受控 AI 工作流：備課、查經、探訪摘要、翻譯校對、事工提醒。"],
        ],
        [1.2, 2.5, 2.8],
    )

    add_heading(doc, "三、已做的初步檢查結果", 1)
    add_bullets(doc, [
        "自動測試 tests/run-all-tests.ps1 目前未通過：test_index_v5_shell.py 仍期待舊的 appendLangGroup 片段，可能是測試未跟上新版 index_v5。",
        "快速斷鏈檢查 scripts/analyze_broken_links_fast.py 目前自身報錯：用文字正則處理 bytes 內容，需先修工具才能產生可信斷鏈報告。",
        "module_audit_report.json 顯示全站模組很多，部分詩歌資料夾含舊網頁與 _vti_cnf 類檔案，後續要區分主線、資料庫、封存、外來資料。",
        "多個 Markdown/規則檔在 PowerShell 顯示為亂碼，初步判斷多半是終端機編碼顯示問題；仍建議統一 UTF-8 保存與檢查。",
    ])

    add_heading(doc, "四、第一階段：先讓網站可維護、可驗證", 1)
    add_numbered(doc, [
        "建立「改任何東西前先備份、改後必測」流程：Git commit、run-all-tests、斷鏈檢查、入口頁人工快速查看。",
        "修正測試腳本與斷鏈腳本，讓測試結果可信；否則後面改 UI 或功能時，不能確定是否真的變好。",
        "為每個主模組建立 MODULE_README：用途、入口頁、可獨立成站時需要哪些檔案、老師可改哪些文字、不可亂動哪些資料。",
        "整理 index_v5：把過長內嵌 CSS/JS 逐步抽到 css/ 與 js/，但每次只移一小段並測試，避免一次大改造成風險。",
        "建立手機版檢查清單：頂欄是否過高、按鈕是否太小、iframe 是否可滾動、離線或 file:// 開啟是否有明確提示。",
    ])

    add_heading(doc, "五、第二階段：改善前端與老師使用體驗", 1)
    add_bullets(doc, [
        "把首頁改成「任務入口」而不只是模組入口：我要備課、我要查經、我要帶小組、我要探訪、我要翻譯教材、我要管理學生。",
        "每個模組有三種模式：老師使用、同工維護、開發者檢查。小白老師只看必要按鈕，維護者才看到資料與測試入口。",
        "統一設計語言：字級、按鈕、顏色、側欄、返回總站、錯誤訊息，避免每個模組像不同網站。",
        "增加低網絡友好設計：清楚標示可離線使用的內容；大量資料延遲載入；必要教材可打包下載。",
        "加入簡明的多語切換與術語表：保留中文/英文核心術語，讓越南、印尼、柬埔寨、老撾等語言維護者容易對照。",
    ])

    add_heading(doc, "六、AI 放進站內的建議路線", 1)
    add_table(
        doc,
        ["AI 功能", "適合先做的版本", "風險控制"],
        [
            ["聖經學習助手", "根據已選經文產生觀察、解釋、應用、問題。", "明確標示 AI 只是輔助；重要教義需老師核對。"],
            ["備課助手", "輸入經文、對象、時間，產生課程大綱、小組問題、禱告回應。", "保留可編輯草稿，不自動當成正式教材。"],
            ["翻譯與校對", "把教材變成目標語言草稿，並保留中英對照。", "建立術語表與人工審核流程。"],
            ["教會事工助手", "探訪記錄整理、代禱事項分類、跟進提醒。", "避免敏感個資外洩；清楚區分本機資料與雲端資料。"],
            ["自動化任務", "定期提醒檢查斷鏈、備份資料、跟進探訪、更新教材。", "先做提醒與報告，不做未審核的自動改檔。"],
        ],
        [1.4, 2.6, 2.5],
    )

    add_heading(doc, "七、自動化任務可怎樣落地", 1)
    add_para(doc, "這裡有兩種不同層次，要分清楚：")
    add_bullets(doc, [
        "Codex app 自動化：在這個工作環境中定期叫我回來檢查、提醒、生成報告，例如每週檢查網站連結與測試結果。",
        "網站內自動化：給老師使用的提醒、待辦、課程排程、探訪跟進，通常需要本機 localStorage、Google Sheet、後端或雲端資料庫配合。",
    ])
    add_para(doc, "建議先從 Codex app 自動化開始，因為成本最低；等網站資料結構穩定後，再做站內提醒與 AI workflow。")

    add_heading(doc, "八、建議的三個發展階段", 1)
    add_table(
        doc,
        ["階段", "目標", "可交付成果"],
        [
            ["0-2 週", "修準檢查工具，整理總入口與小白說明。", "測試修復、斷鏈報告、MODULE_README 範本、小白開站指南。"],
            ["1-2 個月", "改善核心模組 UI 與老師工作流。", "任務式首頁、手機版修正、聖經研讀與 AI 工具整合、老師版教程。"],
            ["3-6 個月", "建立 AI 輔助事工平台。", "備課助手、翻譯校對、探訪跟進、資料同步、定期自動化檢查。"],
        ],
        [1.2, 2.4, 3.3],
    )

    add_heading(doc, "九、我建議你接下來問我的話", 1)
    add_bullets(doc, [
        "請先修好 Bible100 的測試工具，讓 run-all-tests 可以可信地通過。",
        "請幫我做一份小白版 README：怎樣開網站、怎樣找入口、怎樣改文字。",
        "請幫我檢查 index_v5 手機版，有哪些按鈕太小或頁面跑位。",
        "請為 bible_study 模組寫一份 MODULE_README，讓它以後可以獨立成站。",
        "請設計站內 AI 備課助手的第一版流程，但先不要接真 API。",
    ])

    add_heading(doc, "十、結論", 1)
    add_para(doc, "Bible100 已經不是一個空殼，而是一個累積了大量內容、資料和事工想法的平台。下一步最重要的不是立刻加更多功能，而是先建立穩定的維護秩序：入口清楚、測試可信、模組可拆、文件讓小白看得懂。這樣外國山區的老師才可以真正使用、維護，並逐步把 AI 用在聖經學習與教會事工中。")

    doc.core_properties.title = "Bible100 網站改良與發展建議"
    doc.core_properties.subject = "Bible100 static site consulting report"
    doc.core_properties.author = "Codex"
    doc.save(OUT)
    print("DOCX_CREATED")


if __name__ == "__main__":
    build()
